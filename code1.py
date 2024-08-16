import requests
from bs4 import BeautifulSoup
from weasyprint import HTML
import streamlit as st
import tempfile
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from io import BytesIO
from urllib.parse import urljoin

# Function to fetch and parse the webpage
def fetch_webpage(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        return response.text
    except requests.RequestException as e:
        st.error(f"Failed to retrieve the webpage: {e}")
        return None

# Function to fetch external CSS files and include them in the HTML content
def include_css(soup, base_url):
    css_links = soup.find_all("link", rel="stylesheet")
    for link in css_links:
        href = link.get("href")
        if href:
            if not href.startswith("http"):
                href = requests.compat.urljoin(base_url, href)
            try:
                css_response = requests.get(href)
                css_response.raise_for_status()
                style_tag = soup.new_tag("style", type="text/css")
                style_tag.string = css_response.text
                link.replace_with(style_tag)
            except requests.RequestException as e:
                st.warning(f"Failed to retrieve CSS file {href}: {e}")
    return str(soup)

# Function to extract the main content of the webpage
def extract_main_content(soup):
    main_content = soup.find('main') or soup.find('article')
    if not main_content:
        # Try other common main content containers
        main_content = (
            soup.find('div', class_='main-content') or 
            soup.find('div', id='main-content') or 
            soup.find('div', class_='content') or 
            soup.find('div', id='content') or 
            soup.find('div', class_='primary-content')
        )
    if main_content:
        # Remove common sidebar elements
        for sidebar in main_content.find_all(['aside', 'nav', 'header', 'footer']):
            sidebar.decompose()
        return str(main_content)
    else:
        st.warning("Main content not found, using full page content.")
        return str(soup)

# Function to modify the HTML to center-align images and add custom styles
def style_html_content(html_content):
    soup = BeautifulSoup(html_content, "html.parser")

    if not soup.html:
        html_tag = soup.new_tag("html")
        soup.insert(0, html_tag)
    if not soup.head:
        head_tag = soup.new_tag("head")
        soup.html.insert(0, head_tag)

    style_tag = soup.new_tag("style")
    style_tag.string = """
    @import url('https://fonts.googleapis.com/css2?family=Roboto:wght@400;700&display=swap');

    body {
        font-family: 'Roboto', sans-serif;
        line-height: 1.6;
        margin: 0;
        padding: 20px;
    }
    h1, h2, h3, h4, h5, h6 {
        font-weight: 700;
    }
    img {
        display: block;
        margin-left: auto;
        margin-right: auto;
        width: 70%;
        height: auto;
    }
    p {
        text-align: justify;
    }
    """
    soup.head.append(style_tag)

    return str(soup)

# Function to convert the HTML content to PDF
def convert_to_pdf(html_content):
    try:
        html = HTML(string=html_content)
        pdf = html.write_pdf()
        return pdf
    except Exception as e:
        st.error(f"Failed to generate PDF: {e}")
        return None

# Function to convert the HTML content to DOCX
def convert_to_docx(html_content, base_url):
    try:
        soup = BeautifulSoup(html_content, "html.parser")
        doc = Document()

        def process_element(element):
            if element.name == 'h1':
                doc.add_heading(element.get_text(), level=1)
            elif element.name == 'h2':
                doc.add_heading(element.get_text(), level=2)
            elif element.name == 'h3':
                doc.add_heading(element.get_text(), level=3)
            elif element.name == 'p':
                doc.add_paragraph(element.get_text())
            elif element.name == 'img':
                img_url = element.get('src')
                if img_url:
                    # Resolve relative URLs
                    img_url = urljoin(base_url, img_url)
                    try:
                        img_response = requests.get(img_url)
                        img_response.raise_for_status()
                        img_stream = BytesIO(img_response.content)
                        doc.add_picture(img_stream, width=Inches(4))  # Adjust size as needed
                    except requests.RequestException as e:
                        st.warning(f"Failed to retrieve image {img_url}: {e}")

            # Recursively process children of the element
            for child in element.children:
                if isinstance(child, str):
                    doc.add_paragraph(child)
                elif hasattr(child, 'name'):
                    process_element(child)

        # Process the body or the whole HTML if no body tag is found
        body = soup.body if soup.body else soup
        for element in body.children:
            if hasattr(element, 'name'):
                process_element(element)
        
        doc_stream = BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        return doc_stream.read()
    except Exception as e:
        st.error(f"Failed to generate DOCX: {e}")
        return None

# Function to view PDF using PyMuPDF
def view_pdf(pdf_file):
    pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
    num_pages = pdf_document.page_count
    st.write(f"Number of pages: {num_pages}")

    for page_num in range(num_pages):
        page = pdf_document.load_page(page_num)
        pix = page.get_pixmap()
        img = pix.tobytes()
        st.image(img, caption=f"Page {page_num + 1}", use_column_width=True)

# Main function for Streamlit
def main():
    st.title("Webpage to Document Converter")

    num_links = st.number_input("Enter the number of links (1 to 6):", min_value=1, max_value=6, step=1)

    urls = [st.text_input(f"Enter the URL for link {i + 1}:") for i in range(num_links)]

    doc_type = st.selectbox("Select the document format:", options=["PDF", "DOCX"])

    if st.button("Generate Document"):
        combined_html_content = ""

        for url in urls:
            if url:
                html_content = fetch_webpage(url)
                if html_content:
                    soup = BeautifulSoup(html_content, "html.parser")
                    base_url = requests.compat.urljoin(url, '/')
                    html_with_css = include_css(soup, base_url)

                    main_content_html = extract_main_content(BeautifulSoup(html_with_css, "html.parser"))
                    
                    # If main content is not found, use the full page content
                    if main_content_html:
                        styled_html_content = style_html_content(main_content_html)
                    else:
                        styled_html_content = style_html_content(html_with_css)

                    combined_html_content += styled_html_content
                else:
                    st.warning(f"Failed to retrieve the webpage content for URL: {url}")

        if combined_html_content:
            if doc_type == "PDF":
                pdf = convert_to_pdf(combined_html_content)
                if pdf:
                    # Provide download button for PDF
                    st.success("PDF generated successfully!")
                    st.download_button(
                        label="Download PDF",
                        data=pdf,
                        file_name="combined_webpage.pdf",
                        mime="application/pdf"
                    )
                    st.write("Preview of the PDF content:")
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                        temp_file.write(pdf)
                        temp_file.flush()
                        with open(temp_file.name, "rb") as f:
                            view_pdf(f)
            elif doc_type == "DOCX":
                docx_content = convert_to_docx(combined_html_content, base_url)
                if docx_content:
                    # Provide download button for DOCX
                    st.success("DOCX generated successfully!")
                    st.download_button(
                        label="Download DOCX",
                        data=docx_content,
                        file_name="combined_webpage.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        else:
            st.warning("No valid content to generate document.")

if __name__ == "__main__":
    main()
